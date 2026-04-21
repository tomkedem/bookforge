"""
ingest.py

Faithful DOCX ingestion for Hebrew and mixed-direction Word documents.

Goals:
- Preserve paragraph order
- Preserve paragraph layout:
  alignment, direction, left/right indent, first-line indent, hanging indent,
  spacing before/after, line spacing
- Preserve inline formatting (runs)
- Preserve hyperlinks
- Preserve footnotes
- Preserve numbering metadata
- Preserve tables as structured data
- Keep JSON as the source of truth
- Do NOT force HTML wrappers into code blocks
"""

from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

try:
    from docx import Document
except ImportError:
    raise ImportError(
        "\n\n" + "=" * 60 + "\n"
        "חסרה חבילת python-docx!\n"
        "התקן באמצעות: pip install python-docx\n"
        + "=" * 60 + "\n"
    )

# lxml is a transitive dependency of python-docx, but we import it
# explicitly because _extract_images_from_paragraph uses
# etree.tostring() to serialize a single run element for regex
# inspection of its drawing/embed attributes.
from lxml import etree


W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
R_NS = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}"

DEFAULT_FONT_SIZE = 24
SMALL_FONT_THRESHOLD = 18

MONOSPACE_FONTS = {
    "courier",
    "courier new",
    "consolas",
    "monaco",
    "menlo",
    "lucida console",
    "dejavu sans mono",
    "source code pro",
    "fira code",
    "jetbrains mono",
    "inconsolata",
    "roboto mono",
}

LANGUAGE_PATTERNS = {
    "python": [
        r"\bdef\s+\w+\s*\(",
        r"\bclass\s+\w+[:\(]",
        r"\bimport\s+\w+",
        r"\bfrom\s+\w+\s+import",
        r"\bprint\s*\(",
        r"\bif\s+__name__\s*==",
        r"\bself\.",
        r"\bfor\s+\w+\s+in\s+",
        r"\blen\s*\(",
        r"\brange\s*\(",
        r"\breturn\s+",
        r"^\s*#\s*[^\!]",
    ],
    "bash": [
        r"^#!/bin/(ba)?sh",
        r"\$\{?\w+\}?",
        r"\becho\s+",
        r"\bcd\s+",
        r"\bls\b",
        r"\bmkdir\s+",
        r"\brm\s+",
        r"\bcp\s+",
        r"\bmv\s+",
        r"\bcat\s+",
        r"\bgrep\s+",
        r"\bawk\s+",
        r"\bsed\s+",
        r"\bcurl\s+",
        r"\bwget\s+",
        r"\bsudo\s+",
        r"\bapt(?:-get)?\s+",
        r"\bnpm\s+",
        r"\bpip\s+",
        r"\bgit\s+",
        r"\bdocker\s+",
        r"\s+\|\s+",
        r"\s+&&\s+",
    ],
    "javascript": [
        r"\bconst\s+\w+\s*=",
        r"\blet\s+\w+\s*=",
        r"\bvar\s+\w+\s*=",
        r"\bfunction\s+\w+\s*\(",
        r"=>\s*[{\(]?",
        r"\bconsole\.(?:log|error|warn)",
        r"\brequire\s*\(",
        r"\bexport\s+(?:default\s+)?",
        r"\basync\s+function",
        r"\bawait\s+",
        r"\.then\s*\(",
        r"\.forEach\s*\(",
        r"\.map\s*\(",
    ],
    "typescript": [
        r":\s*(?:string|number|boolean|void|any)\b",
        r"\binterface\s+\w+",
        r"\btype\s+\w+\s*=",
        r"<\w+>",
        r"\bas\s+\w+",
    ],
    "json": [
        r'^\s*\{[\s\n]*"',
        r"^\s*\[[\s\n]*[\{\[]",
    ],
    "sql": [
        r"\bSELECT\b",
        r"\bFROM\b",
        r"\bWHERE\b",
        r"\bINSERT\s+INTO\b",
        r"\bUPDATE\b",
        r"\bDELETE\s+FROM\b",
        r"\bCREATE\s+TABLE\b",
    ],
    "html": [
        r"<(?:!DOCTYPE|html|head|body|div|span|p|a|img)",
        r"</\w+>",
    ],
    "css": [
        r"\{\s*[\w-]+\s*:",
        r"\.[a-zA-Z][\w-]*\s*\{",
        r"#[a-zA-Z][\w-]*\s*\{",
    ],
}


def _safe_find(element: Any, xpath: str):
    try:
        return element.find(xpath)
    except Exception:
        return None


def _to_int(value: Optional[str]) -> Optional[int]:
    try:
        return int(value) if value is not None else None
    except Exception:
        return None


def _safe_text(value: Optional[str]) -> str:
    return value or ""


def _safe_paragraph_text(para) -> str:
    try:
        return para.text or ""
    except Exception:
        return ""


def _safe_style_name(para) -> str:
    try:
        style = para.style
        if style is None:
            return "Unknown"
        name = getattr(style, "name", None)
        return name or "Unknown"
    except Exception:
        return "Unknown"


def _build_paragraph_element_map(doc) -> Dict[Any, Any]:
    mapping: Dict[Any, Any] = {}
    try:
        for para in doc.paragraphs:
            mapping[para._element] = para
    except Exception:
        pass
    return mapping


def _extract_document_hyperlinks(doc) -> Dict[str, str]:
    hyperlinks: Dict[str, str] = {}
    try:
        for rel in doc.part.rels.values():
            if "hyperlink" in rel.reltype:
                hyperlinks[rel.rId] = rel._target
    except Exception:
        pass
    return hyperlinks


def _extract_footnotes(doc) -> Dict[str, str]:
    footnotes: Dict[str, str] = {}
    try:
        footnotes_part = doc.part.footnotes_part
        if footnotes_part:
            footnotes_xml = footnotes_part._element
            for fn in footnotes_xml.findall(f".//{W_NS}footnote"):
                fn_id = fn.get(f"{W_NS}id")
                if fn_id and fn_id not in {"0", "-1"}:
                    text = "".join(t.text or "" for t in fn.findall(f".//{W_NS}t")).strip()
                    if text:
                        footnotes[fn_id] = text
    except Exception:
        pass
    return footnotes


def _extract_paragraph_hyperlinks(para, doc_hyperlinks: Dict[str, str]) -> List[Tuple[str, str]]:
    result: List[Tuple[str, str]] = []
    try:
        for hl in para._element.findall(f".//{W_NS}hyperlink"):
            r_id = hl.get(f"{R_NS}id")
            if not r_id or r_id not in doc_hyperlinks:
                continue

            url = doc_hyperlinks[r_id]
            text = "".join(t.text or "" for t in hl.findall(f".//{W_NS}t")).strip()
            if text:
                result.append((text, url))
    except Exception:
        pass
    return result


def _match_run_to_hyperlink(
    run_text: str,
    ordered_hyperlinks: List[Tuple[str, str]],
    cursor: int,
) -> Tuple[Optional[str], int]:
    if not run_text or cursor >= len(ordered_hyperlinks):
        return None, cursor

    link_text, url = ordered_hyperlinks[cursor]

    if run_text == link_text:
        return url, cursor + 1

    if run_text in link_text or link_text in run_text:
        return url, cursor

    return None, cursor


def _detect_code_language(text: str) -> str:
    if not text or not text.strip():
        return ""

    scores: Dict[str, int] = {}

    for lang, patterns in LANGUAGE_PATTERNS.items():
        score = 0
        for pattern in patterns:
            try:
                score += len(re.findall(pattern, text, re.MULTILINE | re.IGNORECASE))
            except Exception:
                continue
        if score > 0:
            scores[lang] = score

    if not scores:
        return ""

    if "typescript" in scores and "javascript" in scores:
        return "typescript" if scores["typescript"] >= 2 else "javascript"

    return max(scores, key=scores.get)


def _empty_format_flags() -> Dict[str, bool]:
    """Return a default format flags dict with all flags set to False."""
    return {
        "bold": False,
        "italic": False,
        "underline": False,
        "code": False,
        "small": False,
        "superscript": False,
        "subscript": False,
    }


def _get_run_format(run) -> Dict[str, bool]:
    """
    Return a dict of formatting flags for this run.

    Each run can have multiple flags set at once (e.g. bold + underline).
    This is a change from the previous behavior where only one format
    could be returned as a string. The dict representation allows
    adjacent runs with identical formatting to be merged correctly,
    which eliminates most of the "broken markdown" artifacts that
    previously required regex cleanup downstream.

    Rules:
    - If the font is monospace, only `code` is set; all other flags
      are False. This matches the previous behavior where code took
      priority, and it's also a Markdown constraint: `code spans`
      cannot contain other inline formatting.
    - superscript and subscript are mutually exclusive by nature;
      if both are somehow set, superscript wins (matches prior behavior).
    """
    flags = _empty_format_flags()

    # Code (monospace font) takes priority and suppresses other flags.
    try:
        font_name = run.font.name
        if font_name and font_name.lower() in MONOSPACE_FONTS:
            flags["code"] = True
            return flags
    except Exception:
        pass

    flags["bold"] = bool(getattr(run, "bold", False))
    flags["italic"] = bool(getattr(run, "italic", False))
    flags["underline"] = bool(getattr(run, "underline", False))

    font_size = None
    try:
        if run.font.size:
            font_size = run.font.size.pt * 2
        else:
            r_pr = _safe_find(run._element, f"{W_NS}rPr")
            if r_pr is not None:
                sz = _safe_find(r_pr, f"{W_NS}sz")
                if sz is not None:
                    font_size = int(sz.get(f"{W_NS}val", DEFAULT_FONT_SIZE))
    except Exception:
        font_size = None

    flags["small"] = bool(font_size and font_size <= SMALL_FONT_THRESHOLD)

    try:
        r_pr = _safe_find(run._element, f"{W_NS}rPr")
        if r_pr is not None:
            vert_align = _safe_find(r_pr, f"{W_NS}vertAlign")
            if vert_align is not None:
                val = vert_align.get(f"{W_NS}val")
                if val == "superscript":
                    flags["superscript"] = True
                elif val == "subscript":
                    flags["subscript"] = True
    except Exception:
        pass

    # Superscript wins over subscript if both are somehow set.
    if flags["superscript"]:
        flags["subscript"] = False

    return flags


def _apply_markdown_format(flags: Dict[str, bool], text: str) -> str:
    """
    Apply markdown/HTML formatting to text based on a flags dict.

    Wrapping order matters for readability of the output. We apply from
    innermost to outermost as follows (outer wrappers are applied last):

      1. code        - innermost, and Markdown disallows other formatting
                       inside a code span, so if code is set, we stop here.
      2. superscript / subscript  (mutually exclusive)
      3. small
      4. underline
      5. italic
      6. bold        - outermost

    The order means a run that is both bold and italic becomes
    `**_text_**` (bold on the outside, italic inside).

    Leading and trailing whitespace inside the text is moved OUTSIDE
    the wrappers before applying them. This matters because Word
    authors often include a trailing space inside a bold run
    (e.g. "AI ") which would otherwise render as "**AI **" in markdown.
    Some renderers tolerate that; others (and assistive tech) treat
    the trailing space as part of the emphasis, which is noisy and
    semantically wrong. Moving the whitespace outside produces
    "**AI** " which renders and reads correctly everywhere.

    If text is empty or whitespace-only, no formatting is applied,
    because `** **` and similar produce malformed markdown.
    """
    if not text:
        return text

    # Code takes over; Markdown code spans cannot contain other formatting.
    if flags.get("code"):
        if "\n" in text:
            return text
        return f"`{text}`"

    # Never wrap whitespace-only text in formatting; it produces broken markdown.
    if not text.strip():
        return text

    # Split out leading/trailing whitespace so wrappers hug just the
    # visible content. We reattach the whitespace after wrapping.
    # lstrip()/rstrip() return the trimmed string; slicing gives us
    # back the removed whitespace parts exactly as they were.
    stripped = text.strip()
    leading_len = len(text) - len(text.lstrip())
    trailing_len = len(text) - len(text.rstrip())
    leading = text[:leading_len]
    trailing = text[len(text) - trailing_len:] if trailing_len else ""

    result = stripped

    if flags.get("superscript"):
        result = f"<sup>{result}</sup>"
    elif flags.get("subscript"):
        result = f"<sub>{result}</sub>"

    if flags.get("small"):
        result = f"<small>{result}</small>"

    if flags.get("underline"):
        result = f"<u>{result}</u>"

    if flags.get("italic"):
        result = f"*{result}*"

    if flags.get("bold"):
        result = f"**{result}**"

    return f"{leading}{result}{trailing}"


def _extract_runs_data(para, doc_hyperlinks: Dict[str, str]) -> List[dict]:
    runs_data: List[dict] = []
    ordered_hyperlinks = _extract_paragraph_hyperlinks(para, doc_hyperlinks)
    hyperlink_cursor = 0

    for run in para.runs:
        run_xml = run._element
        text_parts: List[str] = []

        for child in run_xml:
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag == "t":
                text_parts.append(child.text or "")
            elif tag == "br":
                text_parts.append("\n")
            elif tag == "tab":
                text_parts.append("\t")

        text = "".join(text_parts) if text_parts else (run.text or "")
        if not text:
            continue

        hyperlink, hyperlink_cursor = _match_run_to_hyperlink(
            text, ordered_hyperlinks, hyperlink_cursor
        )

        runs_data.append(
            {
                "text": text,
                "format": _get_run_format(run),
                "hyperlink": hyperlink,
            }
        )

    return runs_data


def _runs_to_markdown(runs_data: List[dict]) -> str:
    """
    Convert a list of run dicts into a single markdown string.

    Runs are merged when their format flags are identical AND their
    hyperlink targets match. This is a stricter criterion than the
    previous code allowed, and it's what eliminates the `**text****text**`
    style artifacts that used to require regex cleanup: when Word
    arbitrarily splits a bold span into two runs with the same
    formatting, they now correctly merge back into one before
    markdown wrapping is applied.
    """
    if not runs_data:
        return ""

    merged: List[dict] = []

    for item in runs_data:
        same_format = (
            merged
            and merged[-1]["format"] == item["format"]
            and merged[-1]["hyperlink"] == item["hyperlink"]
        )
        if same_format:
            merged[-1]["text"] += item["text"]
        else:
            merged.append(dict(item))

    parts: List[str] = []

    for item in merged:
        formatted = _apply_markdown_format(item["format"], item["text"])
        if item["hyperlink"]:
            formatted = f"[{formatted}]({item['hyperlink']})"
        parts.append(formatted)

    return "".join(parts).strip()


def _extract_paragraph_layout(para) -> dict:
    p_pr = _safe_find(para._element, f"{W_NS}pPr")

    layout = {
        "alignment": None,
        "direction": None,
        "left_indent_twips": None,
        "right_indent_twips": None,
        "first_line_indent_twips": None,
        "hanging_indent_twips": None,
        "space_before_twips": None,
        "space_after_twips": None,
        "line_spacing_twips": None,
    }

    if p_pr is None:
        return layout

    jc = _safe_find(p_pr, f"{W_NS}jc")
    if jc is not None:
        val = jc.get(f"{W_NS}val")
        layout["alignment"] = {
            "left": "left",
            "right": "right",
            "center": "center",
            "both": "justify",
            "start": "left",
            "end": "right",
        }.get(val)

    bidi = _safe_find(p_pr, f"{W_NS}bidi")
    layout["direction"] = "rtl" if bidi is not None else "ltr"

    ind = _safe_find(p_pr, f"{W_NS}ind")
    if ind is not None:
        layout["left_indent_twips"] = _to_int(ind.get(f"{W_NS}left") or ind.get(f"{W_NS}start"))
        layout["right_indent_twips"] = _to_int(ind.get(f"{W_NS}right") or ind.get(f"{W_NS}end"))
        layout["first_line_indent_twips"] = _to_int(ind.get(f"{W_NS}firstLine"))
        layout["hanging_indent_twips"] = _to_int(ind.get(f"{W_NS}hanging"))

    spacing = _safe_find(p_pr, f"{W_NS}spacing")
    if spacing is not None:
        layout["space_before_twips"] = _to_int(spacing.get(f"{W_NS}before"))
        layout["space_after_twips"] = _to_int(spacing.get(f"{W_NS}after"))
        layout["line_spacing_twips"] = _to_int(spacing.get(f"{W_NS}line"))

    return layout


def _extract_numbering_formats(doc) -> Dict[Tuple[str, str], Dict[str, Any]]:
    formats: Dict[Tuple[str, str], Dict[str, Any]] = {}

    try:
        numbering_part = doc.part.numbering_part
        if numbering_part is None:
            return formats

        numbering_xml = numbering_part._element
        abstract_nums: Dict[str, Dict[str, Dict[str, Any]]] = {}

        for abstract_num in numbering_xml.findall(f"{W_NS}abstractNum"):
            abstract_id = abstract_num.get(f"{W_NS}abstractNumId")
            if not abstract_id:
                continue

            abstract_nums[abstract_id] = {}

            for lvl in abstract_num.findall(f"{W_NS}lvl"):
                ilvl = lvl.get(f"{W_NS}ilvl", "0")

                num_fmt_elem = _safe_find(lvl, f"{W_NS}numFmt")
                lvl_text_elem = _safe_find(lvl, f"{W_NS}lvlText")
                start_elem = _safe_find(lvl, f"{W_NS}start")

                num_fmt = num_fmt_elem.get(f"{W_NS}val", "decimal") if num_fmt_elem is not None else "decimal"
                lvl_text = lvl_text_elem.get(f"{W_NS}val", "%1.") if lvl_text_elem is not None else "%1."
                start = int(start_elem.get(f"{W_NS}val", "1")) if start_elem is not None else 1

                abstract_nums[abstract_id][ilvl] = {
                    "numFmt": num_fmt,
                    "lvlText": lvl_text,
                    "start": start,
                }

        for num in numbering_xml.findall(f"{W_NS}num"):
            num_id = num.get(f"{W_NS}numId")
            abstract_ref = _safe_find(num, f"{W_NS}abstractNumId")
            if num_id and abstract_ref is not None:
                abstract_id = abstract_ref.get(f"{W_NS}val")
                if abstract_id in abstract_nums:
                    for ilvl, fmt_info in abstract_nums[abstract_id].items():
                        formats[(num_id, ilvl)] = fmt_info

    except Exception:
        pass

    return formats


def _to_roman(num: int) -> str:
    values = [1000, 900, 500, 400, 100, 90, 50, 40, 10, 9, 5, 4, 1]
    symbols = ["M", "CM", "D", "CD", "C", "XC", "L", "XL", "X", "IX", "V", "IV", "I"]
    result = ""
    for i, value in enumerate(values):
        while num >= value:
            result += symbols[i]
            num -= value
    return result


def _to_hebrew_letter(num: int) -> str:
    if num <= 0:
        return str(num)

    ones = ["", "א", "ב", "ג", "ד", "ה", "ו", "ז", "ח", "ט"]
    tens = ["", "י", "כ", "ל", "מ", "נ", "ס", "ע", "פ", "צ"]
    hundreds = ["", "ק", "ר", "ש", "ת"]

    if num < 10:
        return ones[num]
    if num < 100:
        return tens[num // 10] + ones[num % 10]
    if num < 500:
        return hundreds[num // 100] + tens[(num % 100) // 10] + ones[num % 10]
    return str(num)


def _get_bullet_char(lvl_text: str) -> str:
    if not lvl_text:
        return "•"

    bullet_map = {
        "\uf0b7": "•",
        "\uf0a7": "§",
        "\uf076": "◆",
        "\uf0d8": "►",
        "\uf0fc": "✓",
        "\uf06f": "○",
        "\uf06e": "■",
        "\uf0a8": "➢",
        "•": "•",
        "○": "○",
        "●": "●",
        "■": "■",
        "□": "□",
        "◆": "◆",
        "◇": "◇",
        "►": "►",
        "▶": "▶",
        "➢": "➢",
        "➤": "➤",
        "→": "→",
        "✓": "✓",
        "✗": "✗",
        "★": "★",
        "☆": "☆",
        "-": "-",
    }

    if lvl_text in bullet_map:
        return bullet_map[lvl_text]

    for char in lvl_text:
        if char in bullet_map:
            return bullet_map[char]

    return "•"


def _extract_list_data(para, numbering_counters, numbering_formats) -> dict:
    list_data = {
        "num_id": None,
        "ilvl": None,
        "prefix": None,
        "num_format": None,
    }

    try:
        p_pr = _safe_find(para._element, f"{W_NS}pPr")
        if p_pr is None:
            return list_data

        num_pr = _safe_find(p_pr, f"{W_NS}numPr")
        if num_pr is None:
            return list_data

        num_id_elem = _safe_find(num_pr, f"{W_NS}numId")
        ilvl_elem = _safe_find(num_pr, f"{W_NS}ilvl")

        if num_id_elem is None:
            return list_data

        num_id = num_id_elem.get(f"{W_NS}val")
        ilvl = ilvl_elem.get(f"{W_NS}val", "0") if ilvl_elem is not None else "0"

        if num_id == "0":
            return list_data

        key = (num_id, ilvl)
        fmt_info = numbering_formats.get(
            key,
            {"numFmt": "decimal", "lvlText": "%1.", "start": 1},
        )

        if key not in numbering_counters:
            numbering_counters[key] = fmt_info.get("start", 1)
        else:
            numbering_counters[key] += 1

        current_num = numbering_counters[key]
        num_fmt = fmt_info["numFmt"]
        lvl_text = fmt_info["lvlText"]

        if num_fmt == "bullet":
            # Emit a standard markdown bullet marker ("-") instead of
            # the original Unicode glyph from Word (•, ◆, ►, ...).
            #
            # The reason: downstream the text is rendered by Astro's
            # markdown pipeline. When a line starts with "• text",
            # markdown treats it as a plain paragraph that happens to
            # begin with a bullet glyph, producing a flat <p> element.
            # When a line starts with "- text", markdown recognises a
            # list item and produces <ul><li>...</li></ul>. Only the
            # HTML-list form benefits from the reading view's list
            # styles (spacing, hanging indent) and from assistive
            # technology treating it as a list.
            #
            # We keep Word's original bullet shape in list_data for
            # any downstream tooling that wants to inspect it, but
            # the emitted markdown prefix is always "-".
            prefix = "-"
        elif num_fmt == "decimal":
            prefix = lvl_text.replace("%1", str(current_num))
        elif num_fmt == "lowerLetter":
            prefix = lvl_text.replace("%1", chr(ord("a") + (current_num - 1) % 26))
        elif num_fmt == "upperLetter":
            prefix = lvl_text.replace("%1", chr(ord("A") + (current_num - 1) % 26))
        elif num_fmt == "lowerRoman":
            prefix = lvl_text.replace("%1", _to_roman(current_num).lower())
        elif num_fmt == "upperRoman":
            prefix = lvl_text.replace("%1", _to_roman(current_num))
        elif num_fmt in {"hebrew1", "hebrew2"}:
            prefix = lvl_text.replace("%1", _to_hebrew_letter(current_num))
        else:
            prefix = lvl_text.replace("%1", str(current_num))

        list_data["num_id"] = num_id
        list_data["ilvl"] = int(ilvl)
        list_data["prefix"] = prefix
        list_data["num_format"] = num_fmt

    except Exception:
        pass

    return list_data


def _append_footnote_refs(para, markdown_text: str, footnotes: Dict[str, str]) -> str:
    try:
        for fn_ref in para._element.findall(f".//{W_NS}footnoteReference"):
            fn_id = fn_ref.get(f"{W_NS}id")
            if fn_id and fn_id in footnotes:
                markdown_text += f"[^{fn_id}]"
    except Exception:
        pass
    return markdown_text


def _detect_code_style_language(style_name: str) -> Optional[str]:
    """
    If a paragraph is explicitly marked with a Word style named
    "Code <Language>" (e.g., "Code Python", "Code Bash", "Code Yaml"),
    return the language portion in lowercase.

    This supports the recommended authoring workflow: authors create
    paragraph styles per language in Word and apply them to each code
    paragraph. ingest then does not need to guess language from content
    or font heuristics - the author has told it directly.

    Matching is case-insensitive and tolerant of common style-naming
    variants Word allows ("Code Python", "code python", "Code-Python",
    "Code_Python", "CodePython"). Returns None if the style is not a
    Code-* style, so callers can fall through to the font-based
    detection for legacy content.

    The returned language is lowercased and mapped to its canonical
    form where the author used a short name ("yaml" -> "yaml",
    "sh" -> "bash"). This keeps the fenced-code output consistent
    regardless of how the style was named.
    """
    if not style_name:
        return None

    s = style_name.strip().lower()

    # Accept "Code X", "Code-X", "Code_X", or single-token "CodeX".
    lang = None
    for prefix in ("code ", "code-", "code_"):
        if s.startswith(prefix):
            lang = s[len(prefix):].strip()
            break

    # Single-token form ("codepython", "codeyaml") - catches authors
    # who created styles without a separator. Very rare.
    if lang is None and s.startswith("code") and len(s) > 4 and s[4].isalpha():
        lang = s[4:].strip()

    if not lang:
        return None

    # Normalize common language aliases to the form consumers expect
    # in a markdown code fence (```python, ```bash, ```yaml, ...).
    aliases = {
        "py": "python",
        "sh": "bash",
        "shell": "bash",
        "yml": "yaml",
        "js": "javascript",
        "ts": "typescript",
        "md": "markdown",
    }
    return aliases.get(lang, lang)


def _classify_paragraph_block(style_name: str, runs_data: List[dict], markdown_text: str) -> str:
    """
    Classify a paragraph as heading / code / paragraph based on its
    Word paragraph style.

    Rules (in order):
      1. Heading styles ("Heading 1", "Heading 2", or "כותרת") -> heading
      2. Code-* styles ("Code Python", "Code Bash", "Code Yaml",
         etc.) -> code. The language is recovered later from the
         same style name.
      3. Everything else -> paragraph.

    Note on what is NOT a rule:
    We do not promote a paragraph to "code" based on its content
    (keywords like "if", "False", "print") or its run fonts
    (Consolas/Courier on some runs). Content-based detection is
    fragile: Hebrew prose that references a Python identifier
    inline ("name מתפרש כ-False") would otherwise be wrongly
    swept into a code block. Font-based detection is also
    unreliable because Word can insert monospace fonts by auto-
    correct or paste-formatting in places the author did not
    intend.

    The authoring convention is: if something is code, the author
    assigns it a Code-* paragraph style. That single signal is
    authoritative. Inline code spans inside prose are still
    handled through per-run code flags by _apply_markdown_format,
    independently of this classification.
    """
    style_lower = (style_name or "").lower()

    if style_lower.startswith("heading") or "כותרת" in style_lower:
        return "heading"

    if _detect_code_style_language(style_name):
        return "code"

    return "paragraph"


def _cell_has_first_column_style(tc) -> bool:
    """
    Detect whether a table cell is marked as belonging to the table's
    first column via Word's conditional formatting (w:cnfStyle).

    In Word, Table Styles can declare that the first column renders
    in bold (or other emphasis). When that happens, individual runs
    in those cells do NOT carry <w:b> themselves - the bold comes from
    the Table Style through cnfStyle. python-docx's run.bold returns
    None because there is no explicit bold on the run.

    We detect this marker by looking for <w:cnfStyle w:firstColumn="1">
    inside the cell's <w:tcPr>. A similar mechanism exists for
    firstRow, lastRow, lastColumn, and the banded variants, but by
    far the most common pattern in Hebrew technical books is "first
    column is bold" (the label column) and we limit to that.
    """
    try:
        tc_pr = tc.find(f"{W_NS}tcPr")
        if tc_pr is None:
            return False
        cnf = tc_pr.find(f"{W_NS}cnfStyle")
        if cnf is None:
            return False
        first_col = cnf.get(f"{W_NS}firstColumn")
        return first_col == "1"
    except Exception:
        return False


def _force_bold_on_runs(runs: List[dict]) -> List[dict]:
    """
    Return a copy of runs with bold=True forced on every run.

    Used for header rows and conditionally-styled first columns,
    where the bold is inherited from the Table Style rather than
    set directly on the run.

    We don't mutate the input because the same run dicts may be
    used elsewhere in the future. Whitespace-only runs are left
    alone; they would produce "** **" (broken markdown) if wrapped.
    """
    out: List[dict] = []
    for r in runs:
        new_fmt = dict(r.get("format", _empty_format_flags()))
        if r.get("text", "").strip():
            new_fmt["bold"] = True
        out.append({
            "text": r.get("text", ""),
            "format": new_fmt,
            "hyperlink": r.get("hyperlink"),
        })
    return out


def _extract_table_cell_runs(tc, paragraph_map, doc_hyperlinks: Dict[str, str]) -> List[dict]:
    """
    Extract runs from a table cell with full formatting preserved.

    A cell contains one or more paragraphs, and each paragraph contains
    runs. We iterate paragraphs using python-docx's cell.paragraphs
    accessor, then use _extract_runs_data (the same function used for
    body paragraphs) to get runs with format flags.

    Previously this function joined all <w:t> text into one "plain"
    run, which silently stripped every bold/italic/underline from
    tables while the rest of the document preserved formatting.

    Paragraph breaks inside a cell are preserved as newlines between
    the runs of adjacent paragraphs. _build_markdown_table in parse.py
    collapses those newlines to spaces (pipe tables cannot contain
    real line breaks).
    """
    # Build wrapper Paragraph objects from the raw <w:p> elements in
    # the cell. paragraph_map is the map built at ingest start from
    # doc.paragraphs; for table cells we build a small local version
    # since cell paragraphs are not in the top-level paragraph list.
    from docx.text.paragraph import Paragraph

    all_runs: List[dict] = []

    try:
        # Find paragraph elements inside this cell
        cell_paragraphs = tc.findall(f"{W_NS}p")
        if not cell_paragraphs:
            # Fallback: some cells nest paragraphs inside other elements
            cell_paragraphs = tc.findall(f".//{W_NS}p")

        for p_idx, p_elem in enumerate(cell_paragraphs):
            # Wrap the raw element in a Paragraph. The second argument
            # is "parent", which Paragraph uses for style inheritance.
            # Passing None is acceptable here because we only need runs.
            para = Paragraph(p_elem, None)
            para_runs = _extract_runs_data(para, doc_hyperlinks)

            # Insert a newline run between paragraphs inside the cell.
            # This preserves the logical break; parse.py will later
            # collapse newlines to spaces for pipe-table output.
            if p_idx > 0 and para_runs:
                all_runs.append({
                    "text": "\n",
                    "format": _empty_format_flags(),
                    "hyperlink": None,
                })

            all_runs.extend(para_runs)
    except Exception:
        pass

    return all_runs


def _extract_table_data(
    tbl_element,
    doc_index: int,
    paragraph_map: Dict[Any, Any],
    doc_hyperlinks: Dict[str, str],
) -> dict:
    """
    Extract a table into a structured dict.

    Each cell now carries:
    - `runs`: list of run dicts with format flags (bold, italic, etc.)
    - `text`: markdown string built from the runs, with pipe chars
      escaped so they don't break the markdown pipe-table syntax

    Bold inheritance from Table Style:
    Many Word tables style their header row or first column as bold
    via a Table Style rather than by setting <w:b> on individual runs.
    We approximate that here by force-applying bold to:
      1. Every cell in the first row (standard header convention)
      2. Every cell whose <w:cnfStyle w:firstColumn="1"> is set
         (Word's marker that the Table Style's "first column" rule
         applies to this cell)

    This is a heuristic: it will add bold to headers and first-column
    labels that would otherwise look unemphasized in the rendered
    markdown, matching how the table appears in Word. It can produce
    false positives if the author used a Table Style that does NOT
    bold the first column, but that combination is rare in the Hebrew
    technical books this pipeline targets.

    Previously `text` was a plain-text join of <w:t> nodes, which
    dropped every formatting mark from cells.
    """
    rows = []

    try:
        tr_elements = tbl_element.findall(f".//{W_NS}tr")
        for row_idx, tr in enumerate(tr_elements):
            row = []
            is_header_row = (row_idx == 0)
            for tc in tr.findall(f".//{W_NS}tc"):
                runs = _extract_table_cell_runs(tc, paragraph_map, doc_hyperlinks)

                # Apply bold inheritance from Table Style:
                # - Header row gets bold unconditionally
                # - Any cell marked as firstColumn by cnfStyle gets bold
                if is_header_row or _cell_has_first_column_style(tc):
                    runs = _force_bold_on_runs(runs)

                md_text = _runs_to_markdown(runs)
                # Escape pipe characters so they don't break the
                # pipe-table syntax in parse.py. This is safe because
                # markdown does not use raw pipes inside cell content.
                md_text = md_text.replace("|", r"\|")
                row.append(
                    {
                        "text": md_text,
                        "runs": runs,
                    }
                )
            if row:
                rows.append(row)
    except Exception:
        pass

    return {
        "id": f"tbl-{doc_index}",
        "type": "table",
        "doc_index": doc_index,
        "rows": rows,
    }


def _extract_document_metadata(doc, path: Path) -> dict:
    props = None
    try:
        props = doc.core_properties
    except Exception:
        props = None

    title = None
    author = None

    try:
        title = props.title if props and props.title else None
    except Exception:
        title = None

    try:
        author = props.author if props and props.author else None
    except Exception:
        author = None

    if not title:
        title = path.stem

    return {
        "title": title,
        "author": author,
    }


def _extract_images_from_paragraph(p_element, doc_index: int) -> List[dict]:
    """
    Find images embedded in a paragraph element and return positional data.

    Each image gets a record with:
      - doc_index: same counter used by blocks, so parse.py can match
        images to the paragraph they appear in.
      - run_index: position within the paragraph's runs, useful for
        deciding whether the image appears before or after text.
      - rel_id: the relationship id (rId...) pointing to the image
        file in the docx zip; parse.py uses this to locate and save
        the actual bytes.

    Dimensions (width/height in EMU) are read from <wp:extent> so
    parse.py can convert to pixels for the <img> tag.
    """
    images: List[dict] = []

    try:
        # Enumerate runs inside the paragraph
        runs = p_element.findall(f"{W_NS}r")
        for run_idx, r in enumerate(runs):
            try:
                run_xml_bytes = etree.tostring(r)
                run_xml = run_xml_bytes.decode("utf-8") if isinstance(run_xml_bytes, bytes) else run_xml_bytes
            except Exception:
                run_xml = ""

            if not run_xml:
                continue

            if "<w:drawing" not in run_xml and "<w:pict" not in run_xml and "r:embed=" not in run_xml:
                continue

            # Find the relationship id(s) for embedded images
            embeds = re.findall(r'r:embed="(rId\d+)"', run_xml)
            if not embeds:
                continue

            # Try to read display size from <wp:extent cx="..." cy="...">
            w_emu, h_emu = 0, 0
            extents = re.findall(r'<wp:extent\s+cx="(\d+)"\s+cy="(\d+)"', run_xml)
            if extents:
                w_emu, h_emu = int(extents[0][0]), int(extents[0][1])

            for rel_id in embeds:
                images.append({
                    "doc_index": doc_index,
                    "run_index": run_idx,
                    "rel_id": rel_id,
                    "width_emu": w_emu,
                    "height_emu": h_emu,
                    "source": "paragraph",
                })
    except Exception:
        pass

    return images


def _extract_images_from_sdt(sdt_element, doc_index: int) -> List[dict]:
    """
    Find images inside an SDT (Structured Document Tag), which Word
    typically uses for cover-page content. Returns records using the
    same doc_index as the surrounding block loop, so parse.py sees
    a consistent indexing scheme across paragraph and SDT images.
    """
    images: List[dict] = []

    try:
        # Look for <a:blip> anywhere inside the SDT
        blips = sdt_element.xpath(
            ".//a:blip",
            namespaces={"a": "http://schemas.openxmlformats.org/drawingml/2006/main"},
        )
        for blip in blips:
            embed_id = blip.get(
                "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
            )
            if not embed_id:
                continue

            images.append({
                "doc_index": doc_index,
                "run_index": 0,
                "rel_id": embed_id,
                "width_emu": 0,
                "height_emu": 0,
                "source": "sdt",
            })
    except Exception:
        pass

    return images


def _ingest_docx(path: Path, language: str = "he") -> dict:
    doc = Document(path)
    paragraph_map = _build_paragraph_element_map(doc)
    numbering_formats = _extract_numbering_formats(doc)
    numbering_counters: Dict[Tuple[str, str], int] = {}
    doc_hyperlinks = _extract_document_hyperlinks(doc)
    footnotes = _extract_footnotes(doc)
    metadata = _extract_document_metadata(doc, path)

    blocks: List[dict] = []
    # Image records gathered during the same body walk used for blocks.
    # This guarantees that image doc_index values match the block
    # doc_index values; previously parse.py computed image positions
    # separately over doc.paragraphs, which skipped SDT and table
    # elements and drifted out of sync with the block indexing.
    images: List[dict] = []
    body = doc._element.body
    doc_index = 0

    for element in body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

        if tag == "p":
            para = paragraph_map.get(element)
            if para is None:
                doc_index += 1
                continue

            raw_text = _safe_paragraph_text(para)
            style_name = _safe_style_name(para)
            runs_data = _extract_runs_data(para, doc_hyperlinks)
            markdown_text = _runs_to_markdown(runs_data)
            markdown_text = _append_footnote_refs(para, markdown_text, footnotes)
            layout = _extract_paragraph_layout(para)
            list_data = _extract_list_data(para, numbering_counters, numbering_formats)
            block_type = _classify_paragraph_block(style_name, runs_data, raw_text)

            if list_data["prefix"] and block_type != "code":
                indent = "  " * int(list_data["ilvl"] or 0)
                markdown_text = f"{indent}{list_data['prefix']} {markdown_text}"

            # Collect any images embedded in this paragraph. We extract
            # image records *before* deciding whether to skip the block,
            # because a paragraph may contain only an image (no text)
            # and we still need a positional anchor so parse.py can
            # place the image in the output at the right point.
            para_images = _extract_images_from_paragraph(element, doc_index)
            images.extend(para_images)

            if raw_text.strip() or runs_data:
                # For code paragraphs, use the raw text rather than the
                # markdown-formatted text. Two reasons:
                #   1. _apply_markdown_format wraps code runs in
                #      backticks (`foo`), which is correct for inline
                #      code but wrong when the whole paragraph is code
                #      that parse.py will fence with ```...```
                #      - we'd end up with `foo` *inside* a ``` block.
                #   2. _apply_markdown_format strips leading/trailing
                #      whitespace from each run before wrapping. In
                #      Python code, that destroys significant
                #      indentation ("    def foo():" becomes
                #      "def foo():" with the indent sitting outside
                #      the wrapper and then being lost when runs join).
                # raw_text preserves spaces exactly as the author wrote
                # them in Word, which is what a code block needs.
                block_text = raw_text if block_type == "code" else markdown_text

                block = {
                    "id": f"p-{doc_index}",
                    "type": block_type,
                    "text": block_text,
                    "raw_text": raw_text,
                    "style": style_name,
                    "doc_index": doc_index,
                    "runs": runs_data,
                    "layout": layout,
                    "list": list_data,
                }

                if block_type == "code":
                    # Prefer the language declared by the Code-* style
                    # (author-authoritative) over language detection from
                    # content. Fall back to heuristic detection only when
                    # the block was classified as code via the legacy
                    # Consolas-font path and the style provides no hint.
                    style_lang = _detect_code_style_language(style_name)
                    if style_lang:
                        block["language"] = style_lang
                    else:
                        block["language"] = _detect_code_language(raw_text)

                blocks.append(block)
            elif _detect_code_style_language(style_name):
                # Empty paragraph inside a Code-* style. Authors leave
                # blank lines between functions, imports, and logical
                # groups. Without this branch those blank lines would
                # be skipped, collapsing the code block into one dense
                # wall of text. We emit a minimal code block with
                # empty text so parse.py can preserve the blank line
                # when it assembles the fenced block.
                blocks.append({
                    "id": f"p-{doc_index}",
                    "type": "code",
                    "text": "",
                    "raw_text": "",
                    "style": style_name,
                    "doc_index": doc_index,
                    "runs": [],
                    "layout": layout,
                    "list": list_data,
                    "language": _detect_code_style_language(style_name),
                })
            elif para_images:
                # Empty paragraph that holds an image. Emit a minimal
                # block so that parse.py sees a content item at this
                # doc_index and can inject the image inline at the
                # correct position within the chapter. Without this
                # block, the image would be filtered out in to_markdown
                # because no content item matches its para_index.
                blocks.append({
                    "id": f"p-{doc_index}",
                    "type": "image",
                    "text": "",
                    "raw_text": "",
                    "style": style_name,
                    "doc_index": doc_index,
                    "runs": [],
                    "layout": layout,
                    "list": list_data,
                })

            doc_index += 1
            continue

        if tag == "tbl":
            table_block = _extract_table_data(element, doc_index, paragraph_map, doc_hyperlinks)
            if table_block["rows"]:
                blocks.append(table_block)
            doc_index += 1
            continue

        if tag == "sdt":
            # SDT elements carry cover-page images in this document family.
            # We do not currently extract SDT text into blocks (see the
            # --title CLI flag and extract_book_info stub), but we do
            # want to capture cover images.
            images.extend(_extract_images_from_sdt(element, doc_index))
            doc_index += 1
            continue

        doc_index += 1

    return {
        "file": path.name,
        "format": "docx",
        "language": language,
        "metadata": metadata,
        "blocks": blocks,
        "images": images,
        "footnotes": footnotes,
        "total_blocks": len(blocks),
    }


def ingest(file_path: str, language: str = "he") -> dict:
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    suffix = path.suffix.lower()

    if suffix == ".docx":
        return _ingest_docx(path, language=language)

    if suffix == ".pdf":
        raise ValueError(
            "\n" + "=" * 60 + "\n"
            "פורמט PDF לא נתמך.\n"
            "המר את הקובץ ל-Word (.docx) לפני העיבוד.\n"
            + "=" * 60
        )

    raise ValueError(f"Unsupported format: {path.suffix}. Use .docx files only.")


def write_content_structure(structure: dict, output_path: str | Path) -> Path:
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)

    with output.open("w", encoding="utf-8") as f:
        json.dump(structure, f, ensure_ascii=False, indent=2)

    return output


def ingest_and_write_json(
    file_path: str,
    output_path: str | Path,
    language: str = "he",
) -> dict:
    structure = ingest(file_path, language=language)
    write_content_structure(structure, output_path)
    return structure