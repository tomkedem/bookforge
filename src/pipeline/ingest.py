"""
ingest.py

Faithful DOCX ingestion for Hebrew and mixed-direction Word documents.

This module is the orchestrator. It walks the document body once,
delegating the real work to focused sibling modules:

  docx_common  - constants, XML namespaces, tiny safe-access helpers
  docx_code    - Code-* paragraph-style detection, language aliases
  docx_runs    - runs, format flags, markdown wrappers, hyperlinks,
                 footnotes, _runs_to_markdown
  docx_lists   - numbering.xml parsing, per-paragraph list prefixes
  docx_tables  - tables (with Table Style bold inheritance)
  docx_images  - images in paragraphs and in SDT (cover-page) regions

What stays here:
  - _build_paragraph_element_map: one-time index of <w:p> -> Paragraph
  - _extract_paragraph_layout: alignment/indent/spacing per paragraph
  - _extract_document_metadata: title/author from core properties
  - _classify_paragraph_block: heading / code / paragraph decision
  - _ingest_docx: the body-loop that produces the blocks/images list
  - ingest / write_content_structure / ingest_and_write_json: public API

Preserved features (unchanged after the split):
  - Paragraph order
  - Paragraph layout (alignment, direction, indents, spacing)
  - Inline formatting with adjacent-run merging
  - Hyperlinks and footnotes
  - Numbering (bullets, decimal, Roman, Hebrew letters, ...)
  - Structured tables with Table Style bold inheritance
  - JSON as the source of truth
"""

from __future__ import annotations

import json
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

try:
    from docx import Document
except ImportError:
    raise ImportError(
        "\n\n" + "=" * 60 + "\n"
        "חסרה חבילת python-docx!\n"
        "התקן באמצעות: pip install python-docx\n"
        + "=" * 60 + "\n"
    )

from .docx_common import W_NS, _safe_find, _to_int, _safe_paragraph_text, _safe_style_name
from .docx_code import _detect_code_style_language, _detect_code_language
from .docx_runs import (
    _extract_document_hyperlinks,
    _extract_footnotes,
    _extract_runs_data,
    _runs_to_markdown,
    _append_footnote_refs,
)
from .docx_lists import _extract_numbering_formats, _extract_list_data
from .docx_tables import _extract_table_data
from .docx_images import _extract_images_from_paragraph, _extract_images_from_sdt


# ── Small doc-level helpers ────────────────────────────────────────

def _build_paragraph_element_map(doc) -> Dict[Any, Any]:
    """
    Index each paragraph's underlying XML element -> python-docx
    Paragraph wrapper. The body loop iterates raw XML elements for
    ordering, then uses this map to get the wrapper when it needs
    style name, runs, and other python-docx convenience accessors.
    """
    mapping: Dict[Any, Any] = {}
    try:
        for para in doc.paragraphs:
            mapping[para._element] = para
    except Exception:
        pass
    return mapping


def _extract_paragraph_layout(para) -> dict:
    """
    Read a paragraph's visual layout (alignment, text direction,
    indents, spacing). Returns a dict with None for every attribute
    Word didn't set explicitly.

    Units are Word's internal units (twips = 1/20 pt). Callers that
    need pixels do the conversion themselves.
    """
    p_pr = _safe_find(para._element, f"{W_NS}pPr")

    layout = {
        "alignment": None,
        "direction": None,
        "left_indent_twips": None,
        "right_indent_twips": None,
        "first_line_indent_twips": None,
        "hanging_indent_twips": None,
        "space_before_twips": None,
        "space_after_twips": None,
        "line_spacing_twips": None,
    }

    if p_pr is None:
        return layout

    jc = _safe_find(p_pr, f"{W_NS}jc")
    if jc is not None:
        val = jc.get(f"{W_NS}val")
        layout["alignment"] = {
            "left": "left",
            "right": "right",
            "center": "center",
            "both": "justify",
            "start": "left",
            "end": "right",
        }.get(val)

    bidi = _safe_find(p_pr, f"{W_NS}bidi")
    layout["direction"] = "rtl" if bidi is not None else "ltr"

    ind = _safe_find(p_pr, f"{W_NS}ind")
    if ind is not None:
        layout["left_indent_twips"] = _to_int(ind.get(f"{W_NS}left") or ind.get(f"{W_NS}start"))
        layout["right_indent_twips"] = _to_int(ind.get(f"{W_NS}right") or ind.get(f"{W_NS}end"))
        layout["first_line_indent_twips"] = _to_int(ind.get(f"{W_NS}firstLine"))
        layout["hanging_indent_twips"] = _to_int(ind.get(f"{W_NS}hanging"))

    spacing = _safe_find(p_pr, f"{W_NS}spacing")
    if spacing is not None:
        layout["space_before_twips"] = _to_int(spacing.get(f"{W_NS}before"))
        layout["space_after_twips"] = _to_int(spacing.get(f"{W_NS}after"))
        layout["line_spacing_twips"] = _to_int(spacing.get(f"{W_NS}line"))

    return layout


def _extract_document_metadata(doc, path: Path) -> dict:
    """
    Read title/author from the document's core properties. Falls
    back to the file stem when the title property is missing; this
    gives a usable default for files exported without metadata.
    """
    props = None
    try:
        props = doc.core_properties
    except Exception:
        props = None

    title = None
    author = None

    try:
        title = props.title if props and props.title else None
    except Exception:
        title = None

    try:
        author = props.author if props and props.author else None
    except Exception:
        author = None

    if not title:
        title = path.stem

    return {
        "title": title,
        "author": author,
    }


# ── Block classification ───────────────────────────────────────────

def _classify_paragraph_block(style_name: str, runs_data: List[dict], markdown_text: str) -> str:
    """
    Classify a paragraph as heading / code / paragraph based on its
    Word paragraph style.

    Rules (in order):
      1. Heading styles ("Heading 1", "Heading 2", or "כותרת") -> heading
      2. Code-* styles ("Code Python", "Code Bash", "Code Yaml",
         etc.) -> code. The language is recovered later from the
         same style name via _detect_code_style_language.
      3. Everything else -> paragraph.

    Note on what is NOT a rule:
    We do not promote a paragraph to "code" based on its content
    (keywords like "if", "False", "print") or its run fonts
    (Consolas/Courier on some runs). Content-based detection is
    fragile: Hebrew prose that references a Python identifier
    inline ("name מתפרש כ-False") would otherwise be wrongly
    swept into a code block. Font-based detection is also
    unreliable because Word can insert monospace fonts by auto-
    correct or paste-formatting in places the author did not
    intend.

    The authoring convention is: if something is code, the author
    assigns it a Code-* paragraph style. That single signal is
    authoritative. Inline code spans inside prose are still
    handled through per-run code flags by _apply_markdown_format,
    independently of this classification.
    """
    style_lower = (style_name or "").lower()

    if style_lower.startswith("heading") or "כותרת" in style_lower:
        return "heading"

    if _detect_code_style_language(style_name):
        return "code"

    return "paragraph"


# ── Main ingest loop ───────────────────────────────────────────────

def _ingest_docx(path: Path, language: str = "he") -> dict:
    doc = Document(path)
    paragraph_map = _build_paragraph_element_map(doc)
    numbering_formats = _extract_numbering_formats(doc)
    numbering_counters: Dict[Tuple[str, str], int] = {}
    doc_hyperlinks = _extract_document_hyperlinks(doc)
    footnotes = _extract_footnotes(doc)
    metadata = _extract_document_metadata(doc, path)

    blocks: List[dict] = []
    # Image records gathered during the same body walk used for blocks.
    # This guarantees that image doc_index values match the block
    # doc_index values; previously parse.py computed image positions
    # separately over doc.paragraphs, which skipped SDT and table
    # elements and drifted out of sync with the block indexing.
    images: List[dict] = []
    body = doc._element.body
    doc_index = 0

    for element in body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

        if tag == "p":
            para = paragraph_map.get(element)
            if para is None:
                doc_index += 1
                continue

            raw_text = _safe_paragraph_text(para)
            style_name = _safe_style_name(para)
            runs_data = _extract_runs_data(para, doc_hyperlinks)
            markdown_text = _runs_to_markdown(runs_data)
            markdown_text = _append_footnote_refs(para, markdown_text, footnotes)
            layout = _extract_paragraph_layout(para)
            list_data = _extract_list_data(para, numbering_counters, numbering_formats)
            block_type = _classify_paragraph_block(style_name, runs_data, raw_text)

            if list_data["prefix"] and block_type != "code":
                indent = "  " * int(list_data["ilvl"] or 0)
                markdown_text = f"{indent}{list_data['prefix']} {markdown_text}"

            # Collect any images embedded in this paragraph. We extract
            # image records *before* deciding whether to skip the block,
            # because a paragraph may contain only an image (no text)
            # and we still need a positional anchor so parse.py can
            # place the image in the output at the right point.
            para_images = _extract_images_from_paragraph(element, doc_index)
            images.extend(para_images)

            if raw_text.strip() or runs_data:
                # For code paragraphs, use the raw text rather than the
                # markdown-formatted text. Two reasons:
                #   1. _apply_markdown_format wraps code runs in
                #      backticks (`foo`), which is correct for inline
                #      code but wrong when the whole paragraph is code
                #      that parse.py will fence with ```...```
                #      - we'd end up with `foo` *inside* a ``` block.
                #   2. _apply_markdown_format strips leading/trailing
                #      whitespace from each run before wrapping. In
                #      Python code, that destroys significant
                #      indentation ("    def foo():" becomes
                #      "def foo():" with the indent sitting outside
                #      the wrapper and then being lost when runs join).
                # raw_text preserves spaces exactly as the author wrote
                # them in Word, which is what a code block needs.
                block_text = raw_text if block_type == "code" else markdown_text

                block = {
                    "id": f"p-{doc_index}",
                    "type": block_type,
                    "text": block_text,
                    "raw_text": raw_text,
                    "style": style_name,
                    "doc_index": doc_index,
                    "runs": runs_data,
                    "layout": layout,
                    "list": list_data,
                }

                if block_type == "code":
                    # Prefer the language declared by the Code-* style
                    # (author-authoritative) over language detection from
                    # content. Fall back to heuristic detection only when
                    # the style provides no hint (rare in the current
                    # pipeline, which mandates Code-* styles).
                    style_lang = _detect_code_style_language(style_name)
                    if style_lang:
                        block["language"] = style_lang
                    else:
                        block["language"] = _detect_code_language(raw_text)

                blocks.append(block)
            elif _detect_code_style_language(style_name):
                # Empty paragraph inside a Code-* style. Authors leave
                # blank lines between functions, imports, and logical
                # groups. Without this branch those blank lines would
                # be skipped, collapsing the code block into one dense
                # wall of text. We emit a minimal code block with
                # empty text so parse.py can preserve the blank line
                # when it assembles the fenced block.
                blocks.append({
                    "id": f"p-{doc_index}",
                    "type": "code",
                    "text": "",
                    "raw_text": "",
                    "style": style_name,
                    "doc_index": doc_index,
                    "runs": [],
                    "layout": layout,
                    "list": list_data,
                    "language": _detect_code_style_language(style_name),
                })
            elif para_images:
                # Empty paragraph that holds an image. Emit a minimal
                # block so that parse.py sees a content item at this
                # doc_index and can inject the image inline at the
                # correct position within the chapter. Without this
                # block, the image would be filtered out in to_markdown
                # because no content item matches its para_index.
                blocks.append({
                    "id": f"p-{doc_index}",
                    "type": "image",
                    "text": "",
                    "raw_text": "",
                    "style": style_name,
                    "doc_index": doc_index,
                    "runs": [],
                    "layout": layout,
                    "list": list_data,
                })

            doc_index += 1
            continue

        if tag == "tbl":
            table_block = _extract_table_data(element, doc_index, paragraph_map, doc_hyperlinks)
            if table_block["rows"]:
                blocks.append(table_block)
            doc_index += 1
            continue

        if tag == "sdt":
            # SDT elements carry cover-page images in this document family.
            # We do not currently extract SDT text into blocks (see the
            # --title CLI flag on the build driver), but we do want to
            # capture cover images.
            images.extend(_extract_images_from_sdt(element, doc_index))
            doc_index += 1
            continue

        doc_index += 1

    return {
        "file": path.name,
        "format": "docx",
        "language": language,
        "metadata": metadata,
        "blocks": blocks,
        "images": images,
        "footnotes": footnotes,
        "total_blocks": len(blocks),
    }


# ── Public API ─────────────────────────────────────────────────────

def ingest(file_path: str, language: str = "he") -> dict:
    """
    Load a .docx file and return the structured ingestion result.
    PDFs and other formats are rejected here with a clear Hebrew
    message, since the pipeline only supports Word documents.
    """
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    suffix = path.suffix.lower()

    if suffix == ".docx":
        return _ingest_docx(path, language=language)

    if suffix == ".pdf":
        raise ValueError(
            "\n" + "=" * 60 + "\n"
            "פורמט PDF לא נתמך.\n"
            "המר את הקובץ ל-Word (.docx) לפני העיבוד.\n"
            + "=" * 60
        )

    raise ValueError(f"Unsupported format: {path.suffix}. Use .docx files only.")


def write_content_structure(structure: dict, output_path: str | Path) -> Path:
    """Persist an ingestion result as UTF-8 JSON. Returns the path written."""
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)

    with output.open("w", encoding="utf-8") as f:
        json.dump(structure, f, ensure_ascii=False, indent=2)

    return output


def ingest_and_write_json(
    file_path: str,
    output_path: str | Path,
    language: str = "he",
) -> dict:
    """Convenience wrapper that ingests and writes the JSON in one call."""
    structure = ingest(file_path, language=language)
    write_content_structure(structure, output_path)
    return structure